"""Multimodal text extraction from uploaded documents.

Dispatches on content type:
  - PDF            -> pypdf text layer
  - DOCX           -> python-docx paragraphs + tables
  - XLSX           -> openpyxl cell values
  - images         -> Tesseract OCR (pytesseract)
  - text/markdown  -> decoded as-is

Each parser's heavy dependency is imported lazily so the API process doesn't
need the full toolchain — only the Celery worker that runs ingestion does.
Scanned PDFs (no text layer) are flagged for OCR as a hardening follow-up.
"""

from dataclasses import dataclass, field

from app.core.logging import get_logger

log = get_logger(__name__)

IMAGE_TYPES = {"image/png", "image/jpeg", "image/jpg", "image/tiff", "image/bmp"}


@dataclass
class Extraction:
    text: str
    metadata: dict = field(default_factory=dict)


class UnsupportedDocument(Exception):
    pass


def extract(data: bytes, content_type: str, filename: str) -> Extraction:
    ct = (content_type or "").lower()
    name = filename.lower()

    if ct == "application/pdf" or name.endswith(".pdf"):
        return _extract_pdf(data)
    if "wordprocessingml" in ct or name.endswith(".docx"):
        return _extract_docx(data)
    if "spreadsheetml" in ct or name.endswith((".xlsx", ".xlsm")):
        return _extract_xlsx(data)
    if ct in IMAGE_TYPES or name.endswith((".png", ".jpg", ".jpeg", ".tiff", ".bmp")):
        return _extract_image(data)
    if ct.startswith("text/") or name.endswith((".txt", ".md", ".csv")):
        return Extraction(text=data.decode("utf-8", errors="replace"))

    raise UnsupportedDocument(f"Unsupported content type: {content_type}")


def _extract_pdf(data: bytes) -> Extraction:
    import io

    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    text = "\n\n".join(pages).strip()
    meta = {"page_count": len(reader.pages), "needs_ocr": len(text) < 20}
    if meta["needs_ocr"]:
        log.info("pdf_appears_scanned", pages=len(reader.pages))
    return Extraction(text=text, metadata=meta)


def _extract_docx(data: bytes) -> Extraction:
    import io

    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return Extraction(text="\n".join(parts), metadata={"paragraphs": len(parts)})


def _extract_xlsx(data: bytes) -> Extraction:
    import io

    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(f"# Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            values = [str(c) for c in row if c is not None]
            if values:
                parts.append(" | ".join(values))
    return Extraction(text="\n".join(parts), metadata={"sheets": len(wb.worksheets)})


def _extract_image(data: bytes) -> Extraction:
    import io

    import pytesseract
    from PIL import Image

    image = Image.open(io.BytesIO(data))
    text = pytesseract.image_to_string(image)
    return Extraction(
        text=text.strip(),
        metadata={"ocr": True, "width": image.width, "height": image.height},
    )
